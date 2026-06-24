"""
Export utilities — generate PDF and DOCX from a Proposal + ConsultantProfile.
"""

import io
import json


def _safe(text: str) -> str:
    """Escape XML special characters for ReportLab Paragraph."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def generate_pdf(proposal, profile) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_LEFT

    BRAND = colors.HexColor("#4a90d9")
    MUTED = colors.HexColor("#888888")
    DARK = colors.HexColor("#1a1a2e")

    base = getSampleStyleSheet()
    h0 = ParagraphStyle("h0", parent=base["Title"], fontSize=22, textColor=DARK, spaceAfter=4)
    h1_co = ParagraphStyle("h1co", parent=base["Heading1"], fontSize=18, textColor=BRAND, spaceAfter=2)
    sub = ParagraphStyle("sub", parent=base["Normal"], fontSize=11, textColor=MUTED, spaceAfter=10)
    h2 = ParagraphStyle("h2", parent=base["Heading2"], fontSize=14, textColor=DARK, spaceBefore=14, spaceAfter=6)
    body = ParagraphStyle("body", parent=base["Normal"], fontSize=10, leading=15, spaceAfter=5)
    indent = ParagraphStyle("indent", parent=body, leftIndent=16)
    badge_high = ParagraphStyle("badge_h", parent=indent, textColor=colors.HexColor("#cc2200"))
    badge_med = ParagraphStyle("badge_m", parent=indent, textColor=colors.HexColor("#cc7700"))
    badge_low = ParagraphStyle("badge_l", parent=indent, textColor=colors.HexColor("#007700"))

    gap_data = {}
    if proposal.gap_analysis:
        try:
            gap_data = json.loads(proposal.gap_analysis)
        except Exception:
            pass

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=22 * mm, leftMargin=22 * mm,
        topMargin=22 * mm, bottomMargin=22 * mm,
    )

    def hr(color=colors.lightgrey, thickness=0.5):
        return HRFlowable(width="100%", thickness=thickness, color=color, spaceAfter=8)

    story = []

    # ---- Header ----
    story.append(Paragraph("Proposal Package", h0))
    story.append(Paragraph(_safe(proposal.company_name), h1_co))
    if profile:
        story.append(Paragraph(f"Consultant: <b>{_safe(profile.name)}</b>", sub))
    story.append(hr(BRAND, 1))

    # ---- Tailored CV ----
    story.append(Paragraph("Tailored CV", h2))
    story.append(hr())
    for line in proposal.tailored_cv.split("\n"):
        if line.strip():
            story.append(Paragraph(_safe(line), body))
        else:
            story.append(Spacer(1, 4))

    # ---- Bio ----
    story.append(Paragraph("Personalized Bio", h2))
    story.append(hr())
    for line in proposal.bio.split("\n"):
        if line.strip():
            story.append(Paragraph(_safe(line), body))
        else:
            story.append(Spacer(1, 4))

    # ---- Talking Points ----
    story.append(Paragraph("Talking Points", h2))
    story.append(hr())
    for i, point in enumerate(proposal.talking_points, 1):
        story.append(Paragraph(f"{i}.  {_safe(point)}", indent))

    # ---- Gap Analysis ----
    story.append(Paragraph("Gap Analysis", h2))
    story.append(hr())
    if gap_data:
        score = gap_data.get("overall_fit_score", "N/A")
        summary = gap_data.get("overall_fit_summary", "")
        story.append(Paragraph(f"<b>Overall Fit Score: {score} / 10</b>", body))
        if summary:
            story.append(Paragraph(_safe(summary), body))

        strengths = gap_data.get("strengths_to_lead_with", [])
        if strengths:
            story.append(Spacer(1, 6))
            story.append(Paragraph("<b>Lead With These Strengths</b>", body))
            for s in strengths:
                story.append(Paragraph(f"•  {_safe(s)}", indent))

        gaps = gap_data.get("gaps", [])
        if gaps:
            story.append(Spacer(1, 8))
            story.append(Paragraph("<b>Identified Gaps</b>", body))
            severity_style = {"high": badge_high, "medium": badge_med, "low": badge_low}
            for gap in gaps:
                req = gap.get("requirement", "")
                severity = gap.get("severity", "low")
                desc = gap.get("description", "")
                framing = gap.get("framing_suggestion", "")
                mitigation = gap.get("mitigation", "")
                sty = severity_style.get(severity, indent)
                story.append(Paragraph(f"<b>{_safe(req)}</b>  [{severity}]", sty))
                if desc:
                    story.append(Paragraph(_safe(desc), indent))
                if framing:
                    story.append(Paragraph(f"<i>Framing:</i> {_safe(framing)}", indent))
                if mitigation:
                    story.append(Paragraph(f"<i>Mitigation:</i> {_safe(mitigation)}", indent))
                story.append(Spacer(1, 4))
    else:
        story.append(Paragraph("No gap analysis available.", body))

    doc.build(story)
    return buffer.getvalue()


def generate_docx(proposal, profile) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import lxml.etree as etree

    gap_data = {}
    if proposal.gap_analysis:
        try:
            gap_data = json.loads(proposal.gap_analysis)
        except Exception:
            pass

    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    # ---- Header ----
    title = doc.add_heading("Proposal Package", 0)
    co = doc.add_heading(proposal.company_name, 1)
    co.runs[0].font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
    if profile:
        p = doc.add_paragraph()
        run = p.add_run(f"Consultant: {profile.name}")
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(11)

    # ---- Tailored CV ----
    doc.add_heading("Tailored CV", 2)
    doc.add_paragraph(proposal.tailored_cv)

    # ---- Bio ----
    doc.add_heading("Personalized Bio", 2)
    doc.add_paragraph(proposal.bio)

    # ---- Talking Points ----
    doc.add_heading("Talking Points", 2)
    for point in proposal.talking_points:
        doc.add_paragraph(point, style="List Number")

    # ---- Gap Analysis ----
    doc.add_heading("Gap Analysis", 2)
    if gap_data:
        score = gap_data.get("overall_fit_score", "N/A")
        summary = gap_data.get("overall_fit_summary", "")
        p = doc.add_paragraph()
        run = p.add_run(f"Overall Fit Score: {score} / 10")
        run.bold = True
        if summary:
            doc.add_paragraph(summary)

        strengths = gap_data.get("strengths_to_lead_with", [])
        if strengths:
            p = doc.add_paragraph()
            p.add_run("Lead With These Strengths").bold = True
            for s in strengths:
                doc.add_paragraph(s, style="List Bullet")

        gaps = gap_data.get("gaps", [])
        if gaps:
            p = doc.add_paragraph()
            p.add_run("Identified Gaps").bold = True
            for gap in gaps:
                req = gap.get("requirement", "")
                severity = gap.get("severity", "low")
                desc = gap.get("description", "")
                framing = gap.get("framing_suggestion", "")
                mitigation = gap.get("mitigation", "")
                p = doc.add_paragraph()
                run = p.add_run(f"{req}  [{severity}]")
                run.bold = True
                if desc:
                    doc.add_paragraph(desc)
                if framing:
                    doc.add_paragraph(f"Framing: {framing}")
                if mitigation:
                    doc.add_paragraph(f"Mitigation: {mitigation}")
    else:
        doc.add_paragraph("No gap analysis available.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
